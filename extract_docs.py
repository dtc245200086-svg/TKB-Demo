from docx import Document

# Read first document
doc1 = Document(r'c:\Users\phamt\Downloads\TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM TKB.11.docx')
print("=" * 80)
print("DOCUMENT 1: TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM TKB.11")
print("=" * 80)
for p in doc1.paragraphs:
    if p.text.strip():
        print(p.text)

print("\n" * 2)

# Read second document
doc2 = Document(r'c:\Users\phamt\Downloads\TÀI LIỆU BÁO CÁO TỔNG HỢP 22.docx')
print("=" * 80)
print("DOCUMENT 2: TÀI LIỆU BÁO CÁO TỔNG HỢP")
print("=" * 80)
for p in doc2.paragraphs:
    if p.text.strip():
        print(p.text)
